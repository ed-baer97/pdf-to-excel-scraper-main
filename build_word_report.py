import argparse
import json
import re
from pathlib import Path
from datetime import datetime

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt


def _sanitize_filename(s: str) -> str:
    s = " ".join((s or "").split()).strip()
    s = re.sub(r"[<>:\"/\\\\|?*]+", "_", s)
    s = s.strip(" .")
    return s or "report"


def _normalize_text(s: str, replace_yo: bool = False, remove_spaces: bool = False) -> str:
    """
    Normalize text for comparison.
    
    Args:
        s: Input string
        replace_yo: Replace 'ё' with 'е' (default: False)
        remove_spaces: Remove all spaces (default: False)
    
    Returns:
        Normalized lowercase string
    """
    s = (s or "").lower().strip()
    if replace_yo:
        s = s.replace("ё", "е")
    if remove_spaces:
        s = re.sub(r"\s+", "", s)
    else:
        s = re.sub(r"\s+", " ", s).strip()
    return s


def _read_text(path: Path) -> str | None:
    try:
        return path.read_text(encoding="utf-8", errors="ignore").strip() or None
    except Exception:
        return None


def _normalize_name(s: str) -> str:
    s = (s or "").lower()
    s = s.replace("«", "").replace("»", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _resolve_xlsx(path: Path) -> Path:
    """
    If the xlsx path doesn't exist (often due to missing « » in class names),
    try to find the closest match in the same directory by normalized stem.
    """
    if path.exists():
        return path
    parent = path.parent if path.parent else Path(".")
    if not parent.exists():
        return path

    want = _normalize_name(path.stem)
    candidates = list(parent.glob("*.xlsx"))
    if not candidates:
        return path

    # Exact normalized match first
    for c in candidates:
        if _normalize_name(c.stem) == want:
            return c

    # Containment fallback
    for c in candidates:
        cn = _normalize_name(c.stem)
        if want and (want in cn or cn in want):
            return c

    return path


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _replace_in_doc(doc: Document, mapping: dict[str, str]) -> int:
    """
    Best-effort placeholder replacement in paragraphs and table cells.
    Supports: {{KEY}}, <<KEY>>, [KEY], {KEY}.
    Note: If placeholders are split across runs, this may not catch all cases.
    """
    replaced = 0
    patterns = []
    for k, v in mapping.items():
        patterns.append((f"{{{{{k}}}}}", v))
        patterns.append((f"<<{k}>>", v))
        patterns.append((f"[{k}]", v))
        patterns.append((f"{{{k}}}", v))

    for p in _iter_paragraphs(doc):
        txt = p.text
        new_txt = txt
        for pat, val in patterns:
            if pat in new_txt:
                new_txt = new_txt.replace(pat, val)
        if new_txt != txt:
            # Replace whole paragraph text (simple and reliable)
            for r in p.runs:
                r.text = ""
            run = p.add_run(new_txt)
            _apply_font(run)
            replaced += 1
    return replaced


def _fill_goals_table(doc: Document, goals_data: dict) -> bool:
    """
    Find and fill the goals table with structure:
    - First table (index 0)
    - Header row 7 (index 6): "Достигнутые цели" | "Цели, вызвавшие затруднения"
    - Data rows 8-11 (indices 7-10): СОР1, СОР2, СОР3, СОЧ
    
    Russian headers: "Достигнутые цели", "Цели, вызвавшие затруднения"
    Kazakh headers: "Қол жеткізілген мақсаттар", "Қиындық тудырған мақсаттар"
    
    goals_data format:
    {
        "sor1": {"achieved": "...", "difficulties": "..."},
        "sor2": {"achieved": "...", "difficulties": "..."},
        "sor3": {"achieved": "...", "difficulties": "..."},
        "soch": {"achieved": "...", "difficulties": "..."}
    }
    """
    # Keywords for Russian and Kazakh
    # Russian: достигнутые цели, цели вызвавшие затруднения
    # Kazakh: қол жеткізілген мақсаттар, қиындық тудырған мақсаттар
    achieved_keywords = ["достигнут", "достигнутые", "қол жеткіз", "жеткізілген"]
    difficulties_keywords = ["затруднен", "затруднения", "вызвавш", "қиындық", "тудырған"]
    goal_keywords = ["цель", "мақсат"]
    
    # Check if we have at least one table
    if not doc.tables or len(doc.tables) == 0:
        return False
    
    # Try to find goals table - check first table first, then others
    for tbl_idx, tbl in enumerate(doc.tables):
        # Check if table has enough rows (need at least 11 rows: 0-10)
        if len(tbl.rows) < 11:
            continue
        
        # Check multiple possible header rows (6, 5, 4) as structure may vary
        for header_row_idx in [6, 5, 4]:
            if header_row_idx >= len(tbl.rows):
                continue
                
            header_row = tbl.rows[header_row_idx]
            if len(header_row.cells) < 2:
                continue
            
            header_text = [_normalize_text(c.text) for c in header_row.cells]
            header_join = " ".join(header_text)
            
            # Look for goals table headers - Russian or Kazakh
            has_achieved = any(keyword in header_join for keyword in achieved_keywords)
            has_difficulties = any(keyword in header_join for keyword in difficulties_keywords)
            has_goal = any(keyword in header_join for keyword in goal_keywords)
            
            if not ((has_achieved or has_goal) and has_difficulties):
                continue
            
            # Find column indices
            col_achieved = None
            col_difficulties = None
            
            for i, h in enumerate(header_text):
                # Achieved column - has achieved keywords but not difficulties keywords
                if any(keyword in h for keyword in achieved_keywords + goal_keywords):
                    if not any(dk in h for dk in difficulties_keywords):
                        col_achieved = i
                # Difficulties column
                if any(keyword in h for keyword in difficulties_keywords):
                    col_difficulties = i
            
            if col_achieved is None or col_difficulties is None:
                continue
            
            # Fill rows after header (typically 7-10, i.e. header_row_idx+1 to header_row_idx+4)
            row_mapping = [
                (header_row_idx + 1, "sor1"),   # СОР1 / БЖБ1
                (header_row_idx + 2, "sor2"),   # СОР2 / БЖБ2
                (header_row_idx + 3, "sor3"),   # СОР3 / БЖБ3
                (header_row_idx + 4, "soch"),   # СОЧ / ТЖБ
            ]
            
            filled_any = False
            for row_idx, goal_key in row_mapping:
                if row_idx >= len(tbl.rows):
                    continue
                
                if goal_key not in goals_data:
                    continue
                
                cells = tbl.rows[row_idx].cells
                if len(cells) <= max(col_achieved, col_difficulties):
                    continue
                
                goal = goals_data[goal_key]
                
                # Fill achieved goals column
                if col_achieved < len(cells):
                    text = goal.get("achieved", "")
                    if text:
                        _set_cell_text(cells[col_achieved], text)
                
                # Fill difficulties column
                if col_difficulties < len(cells):
                    text = goal.get("difficulties", "")
                    if text:
                        _set_cell_text(cells[col_difficulties], text)
                
                filled_any = True
            
            if filled_any:
                return True
    
    return False


def _fill_goals_table_legacy(doc: Document, goals_data: dict) -> bool:
    """
    Legacy version - kept for reference. Uses fixed first table and row 7.
    """
    # Check if we have at least one table
    if not doc.tables or len(doc.tables) == 0:
        return False
    
    # Use first table (index 0)
    tbl = doc.tables[0]
    
    # Check if table has enough rows (need at least 11 rows: 0-10)
    if len(tbl.rows) < 11:
        return False
    
    # Check row 7 (index 6) for headers
    header_row = tbl.rows[6]
    if len(header_row.cells) < 2:
        return False
    
    header_text = [_normalize_text(c.text) for c in header_row.cells]
    header_join = " ".join(header_text)
    
    # Look for goals table headers
    has_achieved = any(keyword in header_join for keyword in ["достигнут", "достигнутые", "цель"])
    has_difficulties = any(keyword in header_join for keyword in ["затруднен", "затруднения", "вызвавш"])
    
    if not (has_achieved and has_difficulties):
        return False
    
    # Find column indices
    col_achieved = None
    col_difficulties = None
    
    for i, h in enumerate(header_text):
        if any(keyword in h for keyword in ["достигнут", "достигнутые", "цель"]) and "затруднен" not in h:
            col_achieved = i
        if any(keyword in h for keyword in ["затруднен", "затруднения", "вызвавш"]):
            col_difficulties = i
    
    if col_achieved is None or col_difficulties is None:
        return False
    
    # Fill rows 8-11 (indices 7-10) with СОР1, СОР2, СОР3, СОЧ
    row_mapping = [
        (7, "sor1"),   # Row 8 -> СОР1
        (8, "sor2"),   # Row 9 -> СОР2
        (9, "sor3"),   # Row 10 -> СОР3
        (10, "soch"),  # Row 11 -> СОЧ
    ]
    
    filled_any = False
    for row_idx, goal_key in row_mapping:
        if row_idx >= len(tbl.rows):
            continue
        
        if goal_key not in goals_data:
            continue
        
        cells = tbl.rows[row_idx].cells
        if len(cells) <= max(col_achieved, col_difficulties):
            continue
        
        goal = goals_data[goal_key]
        
        # Fill achieved goals column
        if col_achieved < len(cells):
            text = goal.get("achieved", "")
            if text:
                _set_cell_text(cells[col_achieved], text)
        
        # Fill difficulties column
        if col_difficulties < len(cells):
            text = goal.get("difficulties", "")
            if text:
                _set_cell_text(cells[col_difficulties], text)
        
        filled_any = True
    
    return filled_any


def _fill_difficulties_table(doc: Document, goals_data: dict) -> bool:
    """
    Find and fill the third table with structure:
    - Table 3 (index 2)
    - Column 2: Row labels
      - "Перечень затруднений, которые возникли у обучающихся при выполнении заданий"
      - "Причины, указанных выше затруднений у обучающихся при выполнении заданий"
      - "Планируемая коррекционная работа:"
    - Columns 3-5: СОР1, СОР2, СОР3
    - Column 6: СОЧ
    
    goals_data format:
    {
        "sor1": {"difficulties_list": "...", "reasons": "...", "correction": "..."},
        "sor2": {...},
        "sor3": {...},
        "soch": {...}
    }
    """
    # Check if we have at least 3 tables
    if not doc.tables or len(doc.tables) < 3:
        return False
    
    # Use third table (index 2)
    tbl = doc.tables[2]
    
    if len(tbl.rows) < 3:
        return False
    
    # Find rows by text in column 2 (index 1)
    # Support both Russian and Kazakh keywords:
    # Russian: "Перечень затруднений", "Причины затруднений", "Коррекционная работа"
    # Kazakh: "Оқушылардың тапсырмаларды орындау кезіндегі туындаған қиындықтар тізбесі",
    #         "Қиындықтардың себептері", 
    #         "БЖБ талдау қорытындылары бойынша жоспарланған жұмыс"
    row_difficulties = None
    row_reasons = None
    row_correction = None
    
    for r in range(len(tbl.rows)):
        cells = tbl.rows[r].cells
        if len(cells) < 2:
            continue
        
        cell_text = _normalize_text(cells[1].text if len(cells) > 1 else "")
        
        # ВАЖНО: порядок условий имеет значение!
        # Сначала проверяем более специфичные условия, потом общие
        
        # Russian: "Причины, указанных выше затруднений..."
        # Kazakh: "Қиындықтардың себептері"
        if ("причин" in cell_text and "затруднен" in cell_text) or "себептері" in cell_text:
            row_reasons = r
        # Russian: "Перечень затруднений..."
        # Kazakh: "Оқушылардың тапсырмаларды орындау кезіндегі туындаған қиындықтар тізбесі"
        elif ("перечень" in cell_text and "затруднен" in cell_text) or "тізбесі" in cell_text or ("қиындық" in cell_text and "туындаған" in cell_text):
            row_difficulties = r
        # Russian: "Планируемая коррекционная работа"
        # Kazakh: "БЖБ талдау қорытындылары бойынша жоспарланған жұмыс"
        elif "коррекцион" in cell_text or "планируем" in cell_text or "жоспарланған" in cell_text or "талдау қорытындылары" in cell_text:
            row_correction = r
    
    # Логирование для отладки
    print(f"[DEBUG] Table 3 row mapping: difficulties={row_difficulties}, reasons={row_reasons}, correction={row_correction}")
    
    if not any([row_difficulties, row_reasons, row_correction]):
        return False
    
    # Column mapping: 3-5 = СОР1-3 (indices 2-4), 6 = СОЧ (index 5)
    col_mapping = {
        "sor1": 2,  # Column 3
        "sor2": 3,  # Column 4
        "sor3": 4,  # Column 5
        "soch": 5,  # Column 6
    }
    
    filled_any = False
    
    # Fill difficulties row (Перечень затруднений)
    if row_difficulties is not None:
        cells = tbl.rows[row_difficulties].cells
        for goal_key, col_idx in col_mapping.items():
            if goal_key in goals_data and col_idx < len(cells):
                text = goals_data[goal_key].get("difficulties_list", "")
                if text:
                    print(f"[DEBUG] Writing difficulties_list to row={row_difficulties}, col={col_idx} ({goal_key}): {text[:50]}...")
                    _set_cell_text(cells[col_idx], text)
                    filled_any = True
    
    # Fill reasons row (Причины затруднений)
    if row_reasons is not None:
        cells = tbl.rows[row_reasons].cells
        for goal_key, col_idx in col_mapping.items():
            if goal_key in goals_data and col_idx < len(cells):
                text = goals_data[goal_key].get("reasons", "")
                if text:
                    print(f"[DEBUG] Writing reasons to row={row_reasons}, col={col_idx} ({goal_key}): {text[:50]}...")
                    _set_cell_text(cells[col_idx], text)
                    filled_any = True
    
    # Fill correction row (Коррекционная работа)
    if row_correction is not None:
        cells = tbl.rows[row_correction].cells
        for goal_key, col_idx in col_mapping.items():
            if goal_key in goals_data and col_idx < len(cells):
                text = goals_data[goal_key].get("correction", "")
                if text:
                    print(f"[DEBUG] Writing correction to row={row_correction}, col={col_idx} ({goal_key}): {text[:50]}...")
                    _set_cell_text(cells[col_idx], text)
                    filled_any = True
    
    return filled_any


def _apply_font(run) -> None:
    # Enforce Times New Roman 12 for inserted text.
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _set_paragraph_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    run = p.add_run(text)
    _apply_font(run)


def _set_cell_text(cell, text: str) -> None:
    """
    Replace text in the first paragraph of a docx table cell without destroying the cell object.
    """
    if not cell.paragraphs:
        cell.text = str(text)
        # Best-effort apply font to the generated run(s)
        for p in cell.paragraphs:
            for r in p.runs:
                _apply_font(r)
        return
    p = cell.paragraphs[0]
    for r in p.runs:
        r.text = ""
    run = p.add_run(str(text))
    _apply_font(run)


def _extract_subject_from_filename(report_xlsx: Path, class_text: str) -> str:
    """
    Heuristic: report filename is usually "<class> <subject>.xlsx"
    Example: "5 «В» Математика.xlsx" -> subject "Математика"
    """
    stem = report_xlsx.stem
    # 1) Strong regex: strip "<num><letter>" (with/without « ») from the beginning.
    #    Works for both "5В Математика" and "5 «В» Математика".
    m = re.match(
        r"^\s*(\d+)\s*[«\"]?\s*([A-Za-zА-ЯЁӘҒҚҢӨҰҮҺ])\s*[»\"]?\s+(.+?)\s*$",
        stem,
    )
    if m:
        subj = m.group(3).strip()
        return subj or stem

    # 2) If class_text is like "5В", try to remove it from the start in a normalized way.
    cls_norm = _normalize_name(class_text)
    stem_norm = _normalize_name(stem)
    if cls_norm and stem_norm.startswith(cls_norm):
        # Find the first space after the class in the original stem and return the rest
        parts = stem.split(" ", 1)
        if len(parts) == 2:
            return parts[1].strip() or stem
        return stem

    return stem


def _build_goal_from_sheets(sheetnames: list[str]) -> str:
    present = []
    # Normalize names used in our excel generator
    for name in sheetnames:
        n = name.lower()
        if "сор 1" in n or "сор1" in n:
            present.append("СОР 1")
        elif "сор 2" in n or "сор2" in n:
            present.append("СОР 2")
        elif "сор 3" in n or "сор3" in n:
            present.append("СОР 3")
        elif "соч" in n:
            present.append("СОЧ")
    # Keep order
    ordered = []
    for x in ["СОР 1", "СОР 2", "СОР 3", "СОЧ"]:
        if x in present:
            ordered.append(x)
    if not ordered:
        ordered = ["СОР 1", "СОР 2", "СОР 3", "СОЧ"]
    return "Цель: Анализ результатов " + ", ".join(ordered)


def _is_sor_or_soch_sheet(name: str) -> bool:
    n = (name or "").lower()
    # Only include SOR 1..3 and SOCH sheets; exclude "Оценки" and "Формативное..."
    return ("сор" in n) or ("соч" in n)


def _fill_template_lines(doc: Document, *, org: str, period: str | None, subject: str, class_text: str, teacher: str, goal: str, lang: str = "ru") -> int:
    """
    Replace known lines in the provided template by matching their text patterns.
    Supports both Russian (Шаблон.docx) and Kazakh (Шаблон_каз.docx) templates.
    """
    changed = 0
    org_replaced = False
    
    # Language-specific patterns
    is_kazakh = lang.lower() in ("kk", "kaz", "kazakh")
    
    for p in _iter_paragraphs(doc):
        t = p.text.strip()
        tl = t.lower()
        if not t:
            continue

        # 1) Organization name replacement
        # Russian: "наименование организации образования"
        # Kazakh: "білім беру ұйымының атауы"
        if not org_replaced:
            if is_kazakh and "білім беру ұйымының атауы" in tl:
                _set_paragraph_text(p, org)
                changed += 1
                org_replaced = True
                continue
            elif not is_kazakh and "наименование организации образования" in tl:
                _set_paragraph_text(p, org)
                changed += 1
                org_replaced = True
                continue

        # For the rest of replacements we allow multiple substitutions in the SAME paragraph
        new_t = t

        # 2) Period/quarter placeholder
        # Russian: "_ четверть" / "_ полугодие"
        # Kazakh: "мәліметтер _" or "_ тоқсан"
        if period:
            if is_kazakh:
                # Kazakh pattern: "мәліметтер _" - replace underscore with period
                if re.search(r"мәліметтер\s*_+", new_t, flags=re.IGNORECASE):
                    new_t = re.sub(r"(мәліметтер\s*)_+", lambda m: m.group(1) + period, new_t, flags=re.IGNORECASE)
                # Also handle "_ тоқсан" (quarter in Kazakh)
                if re.search(r"_+\s*тоқсан", new_t, flags=re.IGNORECASE):
                    new_t = re.sub(r"_+\s*(тоқсан)", lambda m: period + " " + m.group(1), new_t, flags=re.IGNORECASE)
            else:
                # Russian pattern
                if re.search(r"_+\s*(четверть|полугодие)", new_t, flags=re.IGNORECASE):
                    new_t = re.sub(r"_+\s*(четверть|полугодие)", lambda m: period, new_t, flags=re.IGNORECASE)

        # 3) Subject placeholder
        # Russian: "по предмету _"
        # Kazakh: "_ пәнінен" - subject before "пәнінен"
        if subject and "_" in new_t:
            if is_kazakh:
                # Kazakh: "_ пәнінен" -> "Математика пәнінен"
                if re.search(r"_+\s*пәнінен", new_t, flags=re.IGNORECASE):
                    new_t = re.sub(r"_+\s*(пәнінен)", lambda m: subject + " " + m.group(1), new_t, flags=re.IGNORECASE)
                # Also handle "пәні _" pattern
                elif re.search(r"пәні\s*_+", new_t, flags=re.IGNORECASE):
                    new_t = re.sub(r"(пәні\s*)_+", lambda m: m.group(1) + subject, new_t, flags=re.IGNORECASE)
            else:
                # Russian pattern - use lambda to avoid regex escape issues
                if re.search(r"предмет", new_t, flags=re.IGNORECASE):
                    new_t2 = re.sub(
                        r"((?:по\s+)?предмету\s*)_+",
                        lambda m: m.group(1) + subject,
                        new_t,
                        flags=re.IGNORECASE,
                        count=1,
                    )
                    if new_t2 == new_t:
                        new_t2 = re.sub(
                            r"((?:по\s+)?предмет\s*)_+",
                            lambda m: m.group(1) + subject,
                            new_t,
                            flags=re.IGNORECASE,
                            count=1,
                        )
                    new_t = new_t2

        # 4) Class field
        # Russian: "Класс: _"
        # Kazakh: "Сынып: _"
        if is_kazakh and tl.startswith("сынып:"):
            new_t = f"Сынып: {class_text}"
        elif not is_kazakh and tl.startswith("класс:"):
            new_t = f"Класс: {class_text}"

        # 5) Teacher field
        # Russian: "Педагог: _"
        # Kazakh: "Педагог: _" or "Мұғалім: _"
        if is_kazakh and (tl.startswith("педагог:") or tl.startswith("мұғалім:")):
            prefix = "Мұғалім:" if tl.startswith("мұғалім:") else "Педагог:"
            new_t = f"{prefix} {teacher}"
        elif not is_kazakh and tl.startswith("педагог:"):
            new_t = f"Педагог: {teacher}"

        # 6) Goal field
        # Russian: "Цель: ..."
        # Kazakh: "Мақсат: ..."
        if is_kazakh and tl.startswith("мақсат:"):
            new_t = goal
        elif not is_kazakh and tl.startswith("цель:"):
            new_t = goal

        if new_t != t:
            _set_paragraph_text(p, new_t)
            changed += 1

    return changed


def _fill_existing_analysis_table(doc: Document, blocks: list[dict], lang: str = "ru") -> bool:
    """
    Best-effort: if the template already has a table intended for analysis,
    try to fill it. Otherwise, caller may append a new table.
    
    Supports both Russian and Kazakh templates:
    - Russian: "Класс", "Писали", "Макс", "Кач", "Успев", "5", "4", "3", "2"
    - Kazakh: "БЖБ", "Орындағаны", "Макс балл", "% Сапа", "% Үлгерім", "Ең төмен", "Орта", "Жоғары"
    """
    is_kazakh = lang.lower() in ("kk", "kaz", "kazakh")
    
    def norm(s: str) -> str:
        s = (s or "").lower()
        s = s.replace("ё", "е")
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def norm_kind(s: str) -> str:
        """
        Convert titles like "СОр 1" / "СОР 1" / "СОч" / "БЖБ 1" / "ТЖБ" to canonical keys:
          sor1, sor2, sor3, soch
        Kazakh: БЖБ = СОР (Бөлім бойынша жиынтық бағалау), ТЖБ = СОЧ (Тоқсандық жиынтық бағалау)
        """
        ns = _normalize_text(s, replace_yo=True)
        # Russian: СОЧ, Kazakh: ТЖБ (Тоқсандық жиынтық бағалау)
        if "соч" in ns or "тжб" in ns:
            return "soch"
        # Russian: СОР, Kazakh: БЖБ (Бөлім бойынша жиынтық бағалау)
        m = re.search(r"(?:сор|бжб)\s*([123])", ns)
        if m:
            return f"sor{m.group(1)}"
        return ""

    blocks_by_kind = {}
    for b in blocks:
        k = norm_kind(str(b.get("title", "")))
        if k:
            blocks_by_kind[k] = b

    for tbl in doc.tables:
        if not tbl.rows or len(tbl.rows[0].cells) < 2:
            continue
        header = [_normalize_text(c.text, replace_yo=True) for c in tbl.rows[0].cells]
        header_join = " ".join(header)

        # Check for table markers - Russian or Kazakh
        # Russian: "класс" and "писали"
        # Kazakh: "бжб" and "орындағаны" (or similar structure)
        is_russian_table = "класс" in header_join and "писали" in header_join
        is_kazakh_table = ("бжб" in header_join or "орындағаны" in header_join) and "макс" in header_join
        
        if not is_russian_table and not is_kazakh_table:
            continue

        def find_col(*substrs: str) -> int | None:
            """Find column by any of the given substrings."""
            for i, h in enumerate(header):
                for substr in substrs:
                    if substr in h:
                        return i
            return None

        # Find columns - support both Russian and Kazakh headers
        col_class = find_col("класс", "бжб")  # Class or БЖБ
        col_pisali = find_col("писали", "орындағаны")  # Students count
        col_max = find_col("макс")  # Max points
        col_quality = find_col("кач", "сапа")  # Quality %
        col_success = find_col("успев", "үлгерім")  # Success %
        
        # Grade columns - by level names
        # Высокий/Жоғары = only "5"
        # Средний/Орта = "3" + "4" combined
        # Низкий/Төмен = only "2"
        col_high = find_col("высок", "жоғары", "5")  # High = grade 5
        col_mid = find_col("средн", "орта", "4")     # Medium = grades 3+4
        col_low = find_col("низк", "төмен", "2")     # Low = grade 2
        
        if col_class is None and col_pisali is None:
            continue

        # Fill rows by matching "СОР 1/2/3/СОЧ" or "БЖБ 1/2/3/ТЖБ" in the first column
        filled_kinds: set[str] = set()
        for r in range(1, len(tbl.rows)):
            cells = tbl.rows[r].cells
            # Try first column for the kind identifier
            first_col = col_class if col_class is not None else 0
            class_cell_text = _normalize_text(cells[first_col].text, replace_yo=True)
            k = norm_kind(class_cell_text)
            if not k:
                continue
            # Per user: fill only the FIRST occurrence of each kind.
            if k in filled_kinds:
                continue
            b = blocks_by_kind.get(k)
            if not b:
                continue

            # Fill "Писали/Орындағаны" - number of students
            if col_pisali is not None and col_pisali < len(cells):
                _set_cell_text(cells[col_pisali], str(b.get("students", "")))
            
            # Fill the rest ONLY if corresponding columns exist
            if col_max is not None and col_max < len(cells):
                _set_cell_text(cells[col_max], str(b.get("max", "")))
            if col_quality is not None and col_quality < len(cells):
                _set_cell_text(cells[col_quality], str(b.get("quality", "")))
            if col_success is not None and col_success < len(cells):
                _set_cell_text(cells[col_success], str(b.get("success", "")))
            
            # Grade distribution columns by level:
            # Высокий = only count_5
            if col_high is not None and col_high < len(cells):
                _set_cell_text(cells[col_high], str(b.get("count_5", "")))
            
            # Средний = count_3 + count_4 (combined)
            if col_mid is not None and col_mid < len(cells):
                c3 = b.get("count_3", 0) or 0
                c4 = b.get("count_4", 0) or 0
                try:
                    mid_total = int(c3) + int(c4)
                except (ValueError, TypeError):
                    mid_total = f"{c4}+{c3}" if c3 and c4 else (c4 or c3 or "")
                _set_cell_text(cells[col_mid], str(mid_total))
            
            # Низкий = only count_2
            if col_low is not None and col_low < len(cells):
                _set_cell_text(cells[col_low], str(b.get("count_2", "")))

            filled_kinds.add(k)

        return True

    return False


def _extract_names_from_excel_column(ws, col: str, start_row: int = 8, end_row: int = 39) -> list[str]:
    names: list[str] = []
    for r in range(start_row, end_row + 1):
        v = ws[f"{col}{r}"].value
        if v is None:
            continue
        s = str(v).strip()
        if s:
            names.append(s)
    return names


def _fill_level_table(doc: Document, wb, lang: str = "ru") -> bool:
    """
    Fill the 2nd table:
    Russian: "Уровень", "Баллы", "СОР1", "СОР2", "СОР3", "СОЧ"
    Kazakh: "Деңгей", "Балл", "БЖБ1", "БЖБ2", "БЖБ3", "ТЖБ"
    
    Rows:
    Russian: "Высокий", "Средний", "Низкий"
    Kazakh: "Жоғары", "Орта", "Төмен"

    For each SOR/SOCH sheet in Excel, take names from columns:
      - High   : L
      - Medium : M
      - Low    : N
    and paste into corresponding cell (joined by newlines).
    """
    is_kazakh = lang.lower() in ("kk", "kaz", "kazakh")
    
    # Build mapping from kind -> {level -> names}
    excel_map: dict[str, dict[str, list[str]]] = {}
    for name in wb.sheetnames:
        n = _normalize_text(name, replace_yo=True, remove_spaces=True)
        kind = None
        # Kazakh: ТЖБ = СОЧ, БЖБ = СОР
        if "соч" in n or "тжб" in n:
            kind = "soch"
        else:
            m = re.search(r"(?:сор|бжб)(\d)", n)
            if m:
                kind = f"sor{m.group(1)}"
        if not kind:
            continue
        ws = wb[name]
        excel_map[kind] = {
            "high": _extract_names_from_excel_column(ws, "L"),
            "mid": _extract_names_from_excel_column(ws, "M"),
            "low": _extract_names_from_excel_column(ws, "N"),
        }

    # Find a table with expected headers
    for tbl in doc.tables:
        if not tbl.rows or len(tbl.rows[0].cells) < 4:
            continue
        header_raw = [c.text.strip() for c in tbl.rows[0].cells]
        header = [_normalize_text(x, replace_yo=True, remove_spaces=True) for x in header_raw]
        header_join = " ".join(header)
        
        # Check for table markers - Russian or Kazakh
        is_russian_level_table = "уровень" in header_join and "сор1" in header_join
        is_kazakh_level_table = "деңгей" in header_join.lower() or "бжб1" in header_join
        
        if not is_russian_level_table and not is_kazakh_level_table:
            continue

        # Column indices - support both Russian and Kazakh
        def find_col(*keys: str) -> int | None:
            for key in keys:
                k = _normalize_text(key, replace_yo=True, remove_spaces=True)
                for i, h in enumerate(header):
                    if k == h or k in h:
                        return i
            return None

        col_sor1 = find_col("СОР1", "БЖБ1", "бжб1")
        col_sor2 = find_col("СОР2", "БЖБ2", "бжб2")
        col_sor3 = find_col("СОР3", "БЖБ3", "бжб3")
        col_soch = find_col("СОЧ", "ТЖБ", "тжб")
        col_level = find_col("Уровень", "Деңгей", "деңгей")
        if col_level is None:
            col_level = 0

        # Row mapping by first column - support both Russian and Kazakh
        def row_level_key(text: str) -> str | None:
            t = _normalize_text(text, replace_yo=True, remove_spaces=True).lower()
            # Russian: высокий/средний/низкий
            # Kazakh: жоғары/орта/төмен
            if "высок" in t or "жоғары" in t:
                return "high"
            if "средн" in t or "орта" in t:
                return "mid"
            if "низк" in t or "төмен" in t:
                return "low"
            return None

        filled_any = False
        for r in range(1, len(tbl.rows)):
            cells = tbl.rows[r].cells
            lvl = row_level_key(cells[col_level].text if col_level < len(cells) else "")
            if not lvl:
                continue

            def put(col_idx: int | None, kind: str):
                nonlocal filled_any
                if col_idx is None or col_idx >= len(cells):
                    return
                names = excel_map.get(kind, {}).get(lvl, [])
                if not names:
                    return
                # Append or replace? User asked: "записываем учащихся ФИО".
                # We'll replace the cell text with the names list.
                _set_cell_text(cells[col_idx], "\n".join(names))
                filled_any = True

            put(col_sor1, "sor1")
            put(col_sor2, "sor2")
            put(col_sor3, "sor3")
            put(col_soch, "soch")

        return filled_any

    return False


def _extract_sheet_block(ws) -> dict:
    # Template conventions from your Excel pages:
    org = ws["B1"].value
    class_val = ws["C3"].value
    teacher = ws["C5"].value
    title = ws["C6"].value or ws.title
    students_count = ws["C4"].value

    # If C4 missing, count B8.. until blank
    if not students_count:
        cnt = 0
        for r in range(8, 200):
            v = ws[f"B{r}"].value
            if not v:
                break
            cnt += 1
        students_count = cnt

    # Max points (D8) if present
    max_points = ws["D8"].value

    # Summary (as in template logic)
    quality = ws["J8"].value
    success = ws["K8"].value
    c5 = ws["F41"].value
    c4 = ws["G42"].value
    c3 = ws["H43"].value
    c2 = ws["I44"].value

    return {
        "sheet": ws.title,
        "title": str(title) if title is not None else ws.title,
        "org": str(org) if org is not None else "",
        "class": str(class_val) if class_val is not None else "",
        "teacher": str(teacher) if teacher is not None else "",
        "students": students_count if students_count is not None else "",
        "max": max_points if max_points is not None else "",
        "quality": quality if quality is not None else "",
        "success": success if success is not None else "",
        "count_5": c5 if c5 is not None else "",
        "count_4": c4 if c4 is not None else "",
        "count_3": c3 if c3 is not None else "",
        "count_2": c2 if c2 is not None else "",
    }


def build_word_report(
    template_docx: Path,
    report_xlsx: Path,
    out_dir: Path,
    period_txt: Path | None = None,
    subject_txt: Path | None = None,
    context_json: Path | None = None,
    lang: str = "ru",
) -> Path:
    report_xlsx = _resolve_xlsx(report_xlsx)
    if not report_xlsx.exists():
        raise FileNotFoundError(f"Excel report not found: {report_xlsx}")

    wb = load_workbook(report_xlsx, data_only=True)
    # Only analyze SOR/SOCH sheets (per user: do NOT include formative or grades).
    sheets = [wb[name] for name in wb.sheetnames if _is_sor_or_soch_sheet(name)]
    blocks = [_extract_sheet_block(ws) for ws in sheets]

    # Best guess of global fields
    org = next((b["org"] for b in blocks if b["org"]), "")
    class_val = next((b["class"] for b in blocks if b["class"]), "")
    teacher = next((b["teacher"] for b in blocks if b["teacher"]), "")
    # Prefer subject saved during scraping (context json / subject.txt)
    subject = ""
    try:
        ctx_path = context_json or Path("out/mektep/criteria_context.json")
        if ctx_path and ctx_path.exists():
            ctx = json.loads(ctx_path.read_text(encoding="utf-8"))
            subject = str(ctx.get("subject", "") or "").strip()
    except Exception:
        subject = ""

    if not subject:
        subject = _read_text(subject_txt or Path("out/mektep/subject.txt")) or ""

    if not subject:
        subject = _extract_subject_from_filename(report_xlsx, class_val)
    # Period: prefer context json (period_label), then period.txt
    period = None
    try:
        ctx_path = context_json or Path("out/mektep/criteria_context.json")
        if ctx_path and ctx_path.exists():
            ctx_for_period = json.loads(ctx_path.read_text(encoding="utf-8"))
            period = str(ctx_for_period.get("period_label", "") or "").strip() or None
    except Exception:
        period = None
    if not period:
        period = _read_text(period_txt) if period_txt else None
    goal = _build_goal_from_sheets([ws.title for ws in sheets])

    doc = Document(template_docx)

    # Fill specific template lines requested by user
    _fill_template_lines(
        doc,
        org=org,
        period=period,
        subject=subject,
        class_text=class_val,
        teacher=teacher,
        goal=goal,
        lang=lang,
    )

    mapping = {
        "ORG": org,
        "ORG_NAME": org,
        "CLASS": class_val,
        "TEACHER": teacher,
        "SUBJECT": subject,
        "DATE": datetime.now().strftime("%d.%m.%Y"),
    }
    if period:
        mapping["PERIOD"] = period

    _replace_in_doc(doc, mapping)

    # 7) Table analysis: fill an existing table if template has one; otherwise append.
    filled = _fill_existing_analysis_table(doc, blocks, lang=lang)
    _fill_level_table(doc, wb, lang=lang)
    if not filled:
        doc.add_paragraph("")
        p = doc.add_paragraph("")
        _set_paragraph_text(p, "Сводная таблица по листам отчета:")
        table = doc.add_table(rows=1, cols=8)
        hdr = table.rows[0].cells
        _set_cell_text(hdr[0], "Лист")
        _set_cell_text(hdr[1], "Макс")
        _set_cell_text(hdr[2], "Учащихся")
        _set_cell_text(hdr[3], "Кач-ва %")
        _set_cell_text(hdr[4], "Успев %")
        _set_cell_text(hdr[5], "Высокий")  # Only grade 5
        _set_cell_text(hdr[6], "Средний")  # Grades 3+4 combined
        _set_cell_text(hdr[7], "Низкий")   # Only grade 2

        for b in blocks:
            row = table.add_row().cells
            _set_cell_text(row[0], str(b["title"]))
            _set_cell_text(row[1], str(b["max"]))
            _set_cell_text(row[2], str(b["students"]))
            _set_cell_text(row[3], str(b["quality"]))
            _set_cell_text(row[4], str(b["success"]))
            _set_cell_text(row[5], str(b["count_5"]))  # Высокий = 5
            # Средний = 3 + 4 combined
            c3 = b.get("count_3", 0) or 0
            c4 = b.get("count_4", 0) or 0
            try:
                mid_total = int(c3) + int(c4)
            except (ValueError, TypeError):
                mid_total = f"{c4}+{c3}" if c3 and c4 else (c4 or c3 or "")
            _set_cell_text(row[6], str(mid_total))
            _set_cell_text(row[7], str(b["count_2"]))  # Низкий = 2

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{_sanitize_filename(report_xlsx.stem)}.docx"
    doc.save(out_path)
    return out_path


def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--template", default=None, help="Path to Word template (auto-selects based on --lang if not specified)")
    p.add_argument("--xlsx", required=True, help='Path to Excel report, e.g. "out/mektep/reports/5 В Математика.xlsx"')
    p.add_argument("--outdir", default="out/mektep/reports")
    p.add_argument("--period", default="out/mektep/period.txt")
    p.add_argument("--subjectfile", default="out/mektep/subject.txt")
    p.add_argument("--context", default="out/mektep/criteria_context.json")
    p.add_argument("--lang", default="ru", choices=["ru", "kk"], help="Language: ru (Russian), kk (Kazakh)")
    args = p.parse_args()

    # Auto-select template based on language if not specified
    if args.template:
        template = Path(args.template)
    elif args.lang == "kk":
        template = Path("Шаблон_каз.docx")
        if not template.exists():
            template = Path("Шаблон.docx")
            print(f"Warning: Kazakh template not found, using Russian template")
    else:
        template = Path("Шаблон.docx")

    out_path = build_word_report(
        template_docx=template,
        report_xlsx=_resolve_xlsx(Path(args.xlsx)),
        out_dir=Path(args.outdir),
        period_txt=Path(args.period) if args.period else None,
        subject_txt=Path(args.subjectfile) if args.subjectfile else None,
        context_json=Path(args.context) if args.context else None,
        lang=args.lang,
    )
    print(f"Saved Word report: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

