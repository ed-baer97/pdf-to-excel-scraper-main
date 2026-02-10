"""
Goals Dialog - диалог редактирования целей обучения

Редактирование целей обучения и применение к Word отчетам.
"""
import sys
from pathlib import Path

from PyQt6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QTabWidget, QWidget,
    QLabel, QTextEdit, QPushButton, QListWidget, QMessageBox,
    QDialogButtonBox, QGroupBox, QProgressDialog, QGridLayout,
    QScrollArea, QSizePolicy, QFrame
)
from PyQt6.QtCore import Qt, QSettings
from PyQt6.QtGui import QPixmap

from .reports_manager import ReportsManager
from .translator import get_translator


def _get_logo_path() -> Path:
    """Путь к логотипу (работает в dev и в скомпилированном приложении)."""
    if getattr(sys, 'frozen', False):
        base = Path(sys._MEIPASS)
    else:
        base = Path(__file__).resolve().parent.parent
    return base / "resources" / "img" / "logo_edus_logo_white.png"


class GoalsDialog(QDialog):
    """Диалог редактирования целей"""
    
    def __init__(self, reports_manager: ReportsManager, user_data: dict = None, parent=None):
        super().__init__(parent)
        self.reports_manager = reports_manager
        self.user_data = user_data or {}
        self.settings = QSettings("Mektep", "MektepDesktop")
        self.translator = get_translator()
        saved_lang = self.settings.value("language", "ru")
        self.translator.set_language(saved_lang)
        
        # Словари для хранения виджетов
        self.goals_widgets = {}  # {sor_type: {"achieved": QTextEdit, "difficulties": QTextEdit}}
        self.analysis_widgets = {}  # {(row_type, sor_type): QTextEdit}
        
        self.init_ui()
    
    def init_ui(self):
        """Инициализация интерфейса"""
        self.setWindowTitle(self.translator.tr('goals_title'))
        self.setMinimumSize(1000, 800)
        
        layout = QVBoxLayout(self)
        layout.setSpacing(0)
        
        # === Шапка: логотип + "Цели и задачи преподавания" (как на картинке) ===
        header = QFrame()
        header.setObjectName("goals_header")
        header.setFixedHeight(72)
        header.setStyleSheet("""
            #goals_header {
                background-color: #3d3d3d;
            }
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 1, 15, 1)
        header_layout.setSpacing(12)
        header_layout.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
        
        logo_path = _get_logo_path()
        if logo_path.exists():
            logo_label = QLabel()
            pixmap = QPixmap(str(logo_path))
            logo_label.setPixmap(pixmap.scaled(70, 70, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
            logo_label.setFixedSize(70, 70)
            logo_label.setStyleSheet("background: transparent;")
            header_layout.addWidget(logo_label)
        
        title_label = QLabel(self.translator.tr('goals_title'))
        title_label.setStyleSheet("color: white; font-size: 14pt; font-weight: bold; background: transparent;")
        header_layout.addWidget(title_label)
        header_layout.addStretch()
        
        layout.addWidget(header)
        
        # Скролл для всего содержимого
        layout.setSpacing(15)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        content_layout.setSpacing(15)
        
        # === Секция: Цели обучения (вкладки) ===
        goals_label = "Цели обучения" if self.translator.get_language() == 'ru' else "Оқыту мақсаттары"
        goals_group = QGroupBox(goals_label)
        goals_group_layout = QVBoxLayout()
        
        self.goals_tabs = QTabWidget()
        
        sor_types = ["sor1", "sor2", "sor3", "soch"]
        sor_labels = ["СОР1", "СОР2", "СОР3", "СОЧ"]
        
        for sor_type, sor_label in zip(sor_types, sor_labels):
            tab = self.create_goals_tab(sor_type)
            self.goals_tabs.addTab(tab, sor_label)
        
        goals_group_layout.addWidget(self.goals_tabs)
        goals_group.setLayout(goals_group_layout)
        content_layout.addWidget(goals_group)
        
        # === Кнопка генерации AI ===
        generate_text = "Генерировать анализ (AI)" if self.translator.get_language() == 'ru' else "Талдау жасау (AI)"
        generate_btn = QPushButton(f"✨ {generate_text}")
        generate_btn.setMinimumHeight(35)
        tooltip_text = "Автоматически создать анализ на основе целей текущей вкладки" if self.translator.get_language() == 'ru' else "Ағымдағы қойындының мақсаттары негізінде талдауды автоматты түрде жасау"
        generate_btn.setToolTip(tooltip_text)
        generate_btn.clicked.connect(self.generate_analysis)
        content_layout.addWidget(generate_btn)
        
        # === Секция: Анализ затруднений (таблица 3x4) ===
        analysis_text = "Затруднения и коррекционная работа" if self.translator.get_language() == 'ru' else "Қиындықтар және түзету жұмысы"
        analysis_group = QGroupBox(analysis_text)
        analysis_layout = QGridLayout()
        analysis_layout.setSpacing(8)
        
        # Заголовки столбцов
        headers = ["", "СОР1", "СОР2", "СОР3", "СОЧ"]
        for col, header in enumerate(headers):
            label = QLabel(header)
            label.setStyleSheet("font-weight: bold; padding: 5px;")
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            analysis_layout.addWidget(label, 0, col)
        
        # Строки таблицы
        row_types = [
            ("difficulties_list", "Перечень затруднений, которые возникли у обучающихся при выполнении заданий"),
            ("reasons", "Причины, указанных выше затруднений у обучающихся при выполнении заданий"),
            ("correction", "Планируемая коррекционная работа")
        ]
        
        for row_idx, (row_type, row_label) in enumerate(row_types, start=1):
            # Заголовок строки
            row_header = QLabel(row_label)
            row_header.setWordWrap(True)
            row_header.setMaximumWidth(250)
            row_header.setStyleSheet("padding: 5px; font-size: 11px;")
            analysis_layout.addWidget(row_header, row_idx, 0)
            
            # Ячейки для каждого СОР/СОЧ
            for col_idx, sor_type in enumerate(sor_types, start=1):
                text_edit = QTextEdit()
                text_edit.setMinimumHeight(80)
                text_edit.setMaximumHeight(120)
                text_edit.setPlaceholderText(f"{sor_labels[col_idx-1]}...")
                text_edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
                
                # Сохраняем виджет в словарь
                self.analysis_widgets[(row_type, sor_type)] = text_edit
                analysis_layout.addWidget(text_edit, row_idx, col_idx)
        
        # Установка пропорций столбцов
        analysis_layout.setColumnStretch(0, 2)  # Заголовок
        for i in range(1, 5):
            analysis_layout.setColumnStretch(i, 3)  # Ячейки данных
        
        analysis_group.setLayout(analysis_layout)
        content_layout.addWidget(analysis_group)
        
        # === Секция: Выбор отчетов ===
        reports_group = QGroupBox("Выберите отчеты для применения")
        reports_layout = QVBoxLayout()
        
        self.reports_list = QListWidget()
        self.reports_list.setSelectionMode(QListWidget.SelectionMode.MultiSelection)
        self.reports_list.setMaximumHeight(150)
        self.load_reports_list()
        reports_layout.addWidget(self.reports_list)
        
        reports_group.setLayout(reports_layout)
        content_layout.addWidget(reports_group)
        
        scroll.setWidget(content_widget)
        layout.addWidget(scroll)
        
        # === Кнопки внизу ===
        buttons_layout = QHBoxLayout()
        
        apply_btn = QPushButton("Применить к выбранным отчетам")
        apply_btn.setMinimumHeight(40)
        apply_btn.setStyleSheet("""
            QPushButton {
                background-color: #0d6efd;
                color: white;
                font-weight: bold;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #0b5ed7;
            }
        """)
        apply_btn.clicked.connect(self.apply_goals)
        buttons_layout.addWidget(apply_btn)
        
        close_btn = QPushButton("Закрыть")
        close_btn.setMinimumHeight(40)
        close_btn.clicked.connect(self.close)
        buttons_layout.addWidget(close_btn)
        
        layout.addLayout(buttons_layout)
        
        # Стили
        self.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 1px solid #dee2e6;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 5px;
            }
            QTextEdit {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 5px;
            }
            QTextEdit:focus {
                border: 2px solid #86b7fe;
            }
        """)
    
    def create_goals_tab(self, sor_type: str) -> QWidget:
        """Создать вкладку для целей одного типа оценивания"""
        tab = QWidget()
        layout = QHBoxLayout(tab)
        layout.setSpacing(15)
        
        # Достигнутые цели
        achieved_layout = QVBoxLayout()
        achieved_label = QLabel("Достигнутые цели:")
        achieved_label.setStyleSheet("font-weight: bold;")
        achieved_layout.addWidget(achieved_label)
        
        achieved_text = QTextEdit()
        achieved_text.setPlaceholderText("Введите достигнутые цели обучения...")
        achieved_text.setMinimumHeight(100)
        achieved_layout.addWidget(achieved_text)
        layout.addLayout(achieved_layout)
        
        # Цели с затруднениями
        difficulties_layout = QVBoxLayout()
        difficulties_label = QLabel("Цели, вызвавшие затруднения:")
        difficulties_label.setStyleSheet("font-weight: bold;")
        difficulties_layout.addWidget(difficulties_label)
        
        difficulties_text = QTextEdit()
        difficulties_text.setPlaceholderText("Введите цели, вызвавшие затруднения...")
        difficulties_text.setMinimumHeight(100)
        difficulties_layout.addWidget(difficulties_text)
        layout.addLayout(difficulties_layout)
        
        # Сохраняем виджеты
        self.goals_widgets[sor_type] = {
            "achieved": achieved_text,
            "difficulties": difficulties_text
        }
        
        return tab
    
    def load_reports_list(self):
        """Загрузить список отчетов"""
        self.reports_list.clear()
        reports = self.reports_manager.get_reports()
        
        for report in reports:
            if report.get("word_path") and Path(report["word_path"]).exists():
                item_text = f"{report['class_name']} - {report['subject']} (Четверть {report['period_code']})"
                self.reports_list.addItem(item_text)
                # Сохраняем report в data
                self.reports_list.item(self.reports_list.count() - 1).setData(
                    Qt.ItemDataRole.UserRole, report
                )
    
    def get_current_sor_type(self) -> str:
        """Получить текущий тип СОР по активной вкладке"""
        sor_types = ["sor1", "sor2", "sor3", "soch"]
        return sor_types[self.goals_tabs.currentIndex()]
    
    def generate_analysis(self):
        """Генерация анализа через AI для текущей вкладки"""
        # Получаем AI API ключ и модель из данных пользователя (модель выбирает супер-админ)
        api_key = self.user_data.get("ai_api_key")
        ai_model = self.user_data.get("ai_model") or "qwen-flash-character"
        
        if not api_key:
            QMessageBox.warning(
                self,
                "AI недоступен",
                "AI генерация не настроена для вашей школы.\n\n"
                "Обратитесь к суперадминистратору для добавления AI API ключа."
            )
            return
        
        # Получаем текущий тип СОР
        sor_type = self.get_current_sor_type()
        sor_labels = {"sor1": "СОР1", "sor2": "СОР2", "sor3": "СОР3", "soch": "СОЧ"}
        
        # Получаем данные из текущей вкладки
        achieved = self.goals_widgets[sor_type]["achieved"].toPlainText().strip()
        difficulties = self.goals_widgets[sor_type]["difficulties"].toPlainText().strip()
        
        if not difficulties:
            QMessageBox.warning(
                self,
                "Ошибка",
                f"Заполните поле 'Цели, вызвавшие затруднения' на вкладке {sor_labels[sor_type]} для генерации анализа"
            )
            return
        
        # Показываем курсор ожидания вместо пустого диалога
        from PyQt6.QtWidgets import QApplication
        from PyQt6.QtGui import QCursor
        QApplication.setOverrideCursor(QCursor(Qt.CursorShape.WaitCursor))
        
        try:
            from ai.text_generator import AITextGenerator
            generator = AITextGenerator(api_key, model=ai_model)
            result = generator.generate_analysis(achieved, difficulties)
            
            QApplication.restoreOverrideCursor()
            
            if result.get("success"):
                # Заполняем соответствующие ячейки в таблице анализа
                self.analysis_widgets[("difficulties_list", sor_type)].setText(
                    result.get("difficulties_list", "")
                )
                self.analysis_widgets[("reasons", sor_type)].setText(
                    result.get("reasons", "")
                )
                self.analysis_widgets[("correction", sor_type)].setText(
                    result.get("correction", "")
                )
                
                QMessageBox.information(
                    self,
                    "Успех",
                    f"Анализ для {sor_labels[sor_type]} успешно сгенерирован!"
                )
            else:
                QMessageBox.warning(
                    self,
                    "Ошибка",
                    f"Не удалось сгенерировать анализ:\n{result.get('error', 'Неизвестная ошибка')}"
                )
        
        except ImportError:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(
                self,
                "Ошибка",
                "Модуль AI генерации не найден.\n"
                "Убедитесь, что установлена библиотека openai."
            )
        except Exception as e:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(
                self,
                "Ошибка",
                f"Произошла ошибка при генерации:\n{str(e)}"
            )
    
    def apply_goals(self):
        """Применить цели к выбранным отчетам"""
        # Получаем выбранные отчеты
        selected_items = self.reports_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(
                self,
                "Ошибка",
                "Выберите хотя бы один отчет для применения целей"
            )
            return
        
        # Собираем данные целей
        goals_data = {}
        sor_types = ["sor1", "sor2", "sor3", "soch"]
        
        for sor_type in sor_types:
            goals_data[sor_type] = {
                "achieved": self.goals_widgets[sor_type]["achieved"].toPlainText(),
                "difficulties": self.goals_widgets[sor_type]["difficulties"].toPlainText(),
                "difficulties_list": self.analysis_widgets[("difficulties_list", sor_type)].toPlainText(),
                "reasons": self.analysis_widgets[("reasons", sor_type)].toPlainText(),
                "correction": self.analysis_widgets[("correction", sor_type)].toPlainText(),
            }
        
        # Применяем к каждому отчету
        progress = QProgressDialog("Применение целей...", "Отмена", 0, len(selected_items), self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        
        success_count = 0
        
        for i, item in enumerate(selected_items):
            if progress.wasCanceled():
                break
            
            report = item.data(Qt.ItemDataRole.UserRole)
            progress.setLabelText(f"Обработка: {report['class_name']} - {report['subject']}")
            
            if self._apply_goals_to_report(report, goals_data):
                success_count += 1
            
            progress.setValue(i + 1)
        
        progress.close()
        
        QMessageBox.information(
            self,
            "Готово",
            f"Цели применены к {success_count} из {len(selected_items)} отчетам"
        )
    
    def _apply_goals_to_report(self, report: dict, goals_data: dict) -> bool:
        """Применить цели к одному отчету"""
        try:
            word_path = Path(report["word_path"])
            if not word_path.exists():
                return False
            
            # Импортируем функции заполнения таблиц из build_word_report
            import sys
            from pathlib import Path as PathLib
            
            # Добавляем корень проекта в sys.path для импорта build_word_report
            project_root = PathLib(__file__).resolve().parent.parent.parent
            if str(project_root) not in sys.path:
                sys.path.insert(0, str(project_root))
            
            from build_word_report import _fill_goals_table, _fill_difficulties_table
            from docx import Document
            
            doc = Document(str(word_path))
            
            # Заполняем таблицу целей (первая таблица)
            goals_filled = _fill_goals_table(doc, goals_data)
            
            # Заполняем таблицу затруднений (третья таблица)
            difficulties_filled = _fill_difficulties_table(doc, goals_data)
            
            if goals_filled or difficulties_filled:
                # Сохраняем документ только если что-то заполнили
                doc.save(str(word_path))
                return True
            else:
                # Логируем структуру таблиц для отладки
                print(f"[WARN] Goals/difficulties tables not found in {word_path}")
                print(f"[DEBUG] Found {len(doc.tables)} tables")
                for idx, tbl in enumerate(doc.tables):
                    if tbl.rows:
                        first_row = [c.text.strip()[:30] for c in tbl.rows[0].cells[:3]]
                        print(f"[DEBUG] Table {idx}: {first_row}")
                return False
        
        except ImportError as ie:
            print(f"[ERROR] Import error: {ie}")
            return False
        except Exception as e:
            print(f"[ERROR] Error applying goals to {report['word_path']}: {e}")
            return False
