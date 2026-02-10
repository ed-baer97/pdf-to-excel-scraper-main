import os
import shutil
import threading
import time
import zipfile
import io
from datetime import datetime, timedelta

from flask import Blueprint, current_app, jsonify, redirect, render_template, request, url_for, flash, send_file
from flask_login import login_required, current_user
from pathlib import Path
from sqlalchemy import func

from ..extensions import db
from ..models import Role, ReportFile, ScrapeJob, ScrapeJobStatus, TeacherQuotaUsage, School, User
from ..scraper_runner import run_scrape_job
from ..constants import PERIOD_MAP
from ..redis_utils import ai_rate_limiter


# =============================================================================
# Rate Limiting для AI API (Redis-backed with in-memory fallback)
# =============================================================================

def _check_rate_limit(user_id: int) -> tuple[bool, int]:
    """
    Проверяет rate limit для пользователя.
    Использует Redis если доступен, иначе in-memory.
    
    Returns:
        (allowed: bool, remaining: int) - разрешено ли и сколько осталось
    """
    allowed, reset_in = ai_rate_limiter.is_allowed(str(user_id))
    remaining = ai_rate_limiter.get_remaining(str(user_id))
    return allowed, remaining


def _get_rate_limit_reset_time(user_id: int) -> int:
    """Возвращает секунды до сброса rate limit."""
    allowed, reset_in = ai_rate_limiter.is_allowed(str(user_id))
    return reset_in if not allowed else 0

bp = Blueprint("teacher", __name__, url_prefix="/teacher")


@bp.get("/")
@login_required
def dashboard():
    # Супер-админ может зайти в кабинет учителя (создание отчётов)
    if current_user.role == Role.SCHOOL_ADMIN.value:
        return redirect(url_for("admin.dashboard"))
    files = (
        ReportFile.query.filter_by(teacher_id=current_user.id)
        .order_by(ReportFile.period_code, ReportFile.created_at.desc())
        .limit(200)
        .all()
    )
    
    # Group files by period_code
    files_by_period = {}
    for f in files:
        # Normalize period_code: handle None, empty string, whitespace
        period_code = str(f.period_code).strip() if f.period_code else "unknown"
        
        # Log problematic period_codes for debugging
        if not period_code or period_code == "unknown":
            current_app.logger.warning(
                f"File ID {f.id} ({f.class_name} - {f.subject}) has invalid period_code: '{f.period_code}' "
                f"(raw value: {repr(f.period_code)}, type: {type(f.period_code)})"
            )
        
        if period_code not in files_by_period:
            files_by_period[period_code] = []
        files_by_period[period_code].append(f)
        
        # Debug: log each file's period_code
        if current_app.debug:
            current_app.logger.debug(
                f"File: {f.class_name} - {f.subject}, period_code='{period_code}' "
                f"(raw: {repr(f.period_code)}, type: {type(f.period_code)})"
            )
    
    # Sort periods by code (1, 2, 3, 4)
    sorted_periods = sorted(files_by_period.keys(), key=lambda x: int(x) if x.isdigit() else 999)
    
    # Calculate total files count
    total_files = sum(len(files_list) for files_list in files_by_period.values())
    
    # Get latest job for progress bar
    latest_job = (
        ScrapeJob.query.filter_by(teacher_id=current_user.id)
        .order_by(ScrapeJob.created_at.desc())
        .first()
    )
    # Debug: log file count and period distribution
    period_counts = {period: len(files_by_period.get(period, [])) for period in ["1", "2", "3", "4"]}
    # Also log any unexpected period codes
    unexpected_periods = {k: len(v) for k, v in files_by_period.items() if k not in ["1", "2", "3", "4", "unknown"]}
    
    current_app.logger.info(
        f"Teacher {current_user.id} ({current_user.username}): "
        f"found {total_files} files total. "
        f"By period: {period_counts}. "
        f"Sorted periods: {sorted_periods}"
        + (f" Unexpected periods: {unexpected_periods}" if unexpected_periods else "")
    )
    return render_template("teacher/dashboard.html", files_by_period=files_by_period, sorted_periods=sorted_periods, period_map=PERIOD_MAP, total_files=total_files, latest_job=latest_job)


@bp.post("/scrape")
@login_required
def start_scrape():
    if current_user.role != Role.TEACHER.value:
        return redirect(url_for("teacher.dashboard"))

    mektep_login = request.form.get("mektep_login", "").strip()
    mektep_password = request.form.get("mektep_password", "")
    period_code = request.form.get("period_code", "2").strip() or "2"
    lang = request.form.get("lang", "ru").strip() or "ru"
    # school_index теперь не передается из формы - будет динамический выбор если нужен
    school_index = ""

    if not mektep_login or not mektep_password:
        flash("Введите логин/пароль от mektep.edu.kz.", "danger")
        return redirect(url_for("teacher.dashboard"))

    school: School | None = current_user.school
    if not school or not school.is_active:
        flash("Доступ школы закрыт.", "danger")
        return redirect(url_for("teacher.dashboard"))

    quota = int(school.reports_quota_per_period or 0)
    usage = TeacherQuotaUsage.query.filter_by(teacher_id=current_user.id, period_code=period_code).first()
    used = int(usage.used_reports) if usage else 0
    remaining = max(0, quota - used)
    if remaining <= 0:
        flash(f"Лимит успешных скрапов на эту четверть исчерпан (квота {quota}).", "danger")
        return redirect(url_for("teacher.dashboard"))

    # Ensure teacher has per-school filesystem sequence (teacher_1, teacher_2, ...)
    teacher: User | None = db.session.get(User, current_user.id)
    if teacher and teacher.fs_teacher_seq is None:
        max_seq = (
            db.session.query(func.max(User.fs_teacher_seq))
            .filter(User.school_id == school.id, User.role == Role.TEACHER.value)
            .scalar()
        )
        teacher.fs_teacher_seq = int(max_seq or 0) + 1
        db.session.commit()

    job = ScrapeJob(
        school_id=school.id,
        teacher_id=current_user.id,
        period_code=period_code,
        lang=lang,
        status=ScrapeJobStatus.QUEUED.value,
    )
    # Assign per-(school, teacher) sequential number for filesystem paths (job_1, job_2, ...)
    max_job_seq = (
        db.session.query(func.max(ScrapeJob.fs_job_seq))
        .filter(ScrapeJob.school_id == school.id, ScrapeJob.teacher_id == current_user.id)
        .scalar()
    )
    job.fs_job_seq = int(max_job_seq or 0) + 1
    db.session.add(job)
    db.session.commit()

    # Check if Celery is enabled for background jobs
    use_celery = current_app.config.get("USE_CELERY", False)
    app_obj = current_app._get_current_object()
    
    if use_celery:
        try:
            from ..tasks import run_scrape_task
            from ..scraper_runner import _resolve_upload_root
            
            base_upload_dir = _resolve_upload_root(app_obj)
            job_dir = base_upload_dir / f"school_{school.id}" / f"teacher_{current_user.fs_teacher_seq or current_user.id}" / f"job_{job.fs_job_seq}"
            job_dir.mkdir(parents=True, exist_ok=True)
            
            job.output_dir = str(job_dir)
            db.session.commit()
            
            task = run_scrape_task.delay(
                job_id=job.id,
                login=mektep_login,
                password=mektep_password,
                output_dir=str(job_dir),
                period_code=period_code,
                lang=lang,
                school_index=school_index,
            )
            job.celery_task_id = task.id
            db.session.commit()
            current_app.logger.info(f"Job {job.id} queued as Celery task {task.id}")
        except Exception as e:
            current_app.logger.error(f"Celery failed, using thread: {e}")
            use_celery = False
    
    if not use_celery:
        t = threading.Thread(
            target=run_scrape_job,
            kwargs={
                "app": app_obj,
                "job_id": job.id,
                "mektep_login": mektep_login,
                "mektep_password": mektep_password,
                "period_code": period_code,
                "lang": lang,
                "school_index": school_index,
                "limit": 0,  # квота считает успешные скрапы, лимит отчётов за запуск не задаём  # “первые N отчетов”
                "headless": True,
            },
            daemon=True,
        )
        t.start()

    flash(f"Задача запущена. Осталось успешных скрапов в четверти: {remaining}.", "success")
    return redirect(url_for("teacher.dashboard"))


@bp.get("/jobs/latest/status")
@login_required
def get_latest_job_status():
    """AJAX endpoint: return status of latest job."""
    if current_user.role != Role.TEACHER.value:
        return jsonify({"error": "Unauthorized"}), 403
    
    latest_job = (
        ScrapeJob.query.filter_by(teacher_id=current_user.id)
        .order_by(ScrapeJob.created_at.desc())
        .first()
    )
    
    if not latest_job:
        return jsonify({"status": "none", "message": "Нет активных задач"})
    
    # Автоматический сброс зависших задач (>10 минут без прогресса)
    if latest_job.status in (ScrapeJobStatus.RUNNING.value, ScrapeJobStatus.QUEUED.value):
        if latest_job.created_at:
            elapsed = (datetime.utcnow() - latest_job.created_at).total_seconds()
            stuck_threshold = 10 * 60  # 10 минут
            
            # Если задача висит >10 минут и нет прогресса
            if elapsed > stuck_threshold and (not latest_job.progress_percent or latest_job.progress_percent < 10):
                current_app.logger.warning(
                    f"Auto-marking job {latest_job.id} as failed (stuck for {elapsed/60:.1f} min)"
                )
                latest_job.status = ScrapeJobStatus.FAILED.value
                latest_job.error = f"Задача зависла (>{stuck_threshold//60} мин без прогресса). Попробуйте снова."
                latest_job.finished_at = datetime.utcnow()
                db.session.commit()
    
    return jsonify({
        "status": latest_job.status,
        "job_id": latest_job.id,
        "error": latest_job.error,
        "progress_percent": latest_job.progress_percent or 0,
        "progress_message": latest_job.progress_message,
        "total_reports": latest_job.total_reports,
        "processed_reports": latest_job.processed_reports or 0,
        "created_at": latest_job.created_at.isoformat() if latest_job.created_at else None,
        "finished_at": latest_job.finished_at.isoformat() if latest_job.finished_at else None,
    })


def _can_access_file(rf: ReportFile) -> bool:
    if current_user.role == Role.SUPERADMIN.value:
        return True
    if current_user.role == Role.SCHOOL_ADMIN.value:
        return rf.school_id == current_user.school_id
    return rf.teacher_id == current_user.id


@bp.get("/files/<int:file_id>/excel")
@login_required
def download_excel(file_id: int):
    rf = db.session.get(ReportFile, file_id)
    if not rf or not rf.excel_path or not _can_access_file(rf):
        flash("Файл не найден или доступ запрещен.", "danger")
        return redirect(url_for("teacher.dashboard"))
    p = Path(rf.excel_path)
    if not p.exists():
        flash(f"Файл не найден: {p}", "danger")
        return redirect(url_for("teacher.dashboard"))
    return send_file(str(p), as_attachment=True, download_name=p.name)


@bp.get("/files/<int:file_id>/word")
@login_required
def download_word(file_id: int):
    rf = db.session.get(ReportFile, file_id)
    if not rf or not rf.word_path or not _can_access_file(rf):
        flash("Файл не найден или доступ запрещен.", "danger")
        return redirect(url_for("teacher.dashboard"))
    p = Path(rf.word_path)
    if not p.exists():
        flash(f"Файл не найден: {p}", "danger")
        return redirect(url_for("teacher.dashboard"))
    return send_file(str(p), as_attachment=True, download_name=p.name)


@bp.get("/files/download-all")
@login_required
def download_all_files():
    """Download all report files for the selected period as a ZIP archive."""
    if current_user.role != Role.TEACHER.value:
        flash("Доступ запрещен.", "danger")
        return redirect(url_for("teacher.dashboard"))
    
    period_code = request.args.get("period_code", "").strip()
    if not period_code:
        flash("Не указана четверть.", "danger")
        return redirect(url_for("teacher.dashboard"))
    
    # Get all files for this teacher and period
    files = (
        ReportFile.query.filter_by(teacher_id=current_user.id, period_code=period_code)
        .order_by(ReportFile.class_name, ReportFile.subject)
        .all()
    )
    
    if not files:
        flash(f"Нет отчетов для четверти {period_code}.", "info")
        return redirect(url_for("teacher.dashboard"))
    
    # Create ZIP archive in memory
    zip_buffer = io.BytesIO()
    
    try:
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            added_count = 0
            
            for rf in files:
                # Add Excel file if exists
                if rf.excel_path:
                    excel_path = Path(rf.excel_path)
                    if excel_path.exists():
                        # Create a safe filename
                        safe_name = f"{rf.class_name}_{rf.subject}".replace("/", "_").replace("\\", "_")
                        excel_name = f"{safe_name}.xlsx"
                        zip_file.write(str(excel_path), excel_name)
                        added_count += 1
                
                # Add Word file if exists
                if rf.word_path:
                    word_path = Path(rf.word_path)
                    if word_path.exists():
                        # Create a safe filename
                        safe_name = f"{rf.class_name}_{rf.subject}".replace("/", "_").replace("\\", "_")
                        word_name = f"{safe_name}.docx"
                        zip_file.write(str(word_path), word_name)
                        added_count += 1
        
        if added_count == 0:
            flash("Не найдено файлов для скачивания.", "warning")
            return redirect(url_for("teacher.dashboard"))
        
        # Prepare response
        zip_buffer.seek(0)
        period_label = PERIOD_MAP.get(period_code, f"Четверть {period_code}")
        zip_filename = f"Отчеты_{period_label}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=zip_filename
        )
    
    except Exception as e:
        current_app.logger.error(f"Error creating ZIP archive: {e}", exc_info=True)
        flash(f"Ошибка при создании архива: {str(e)}", "danger")
        return redirect(url_for("teacher.dashboard"))


@bp.post("/jobs/<int:job_id>/cancel")
@login_required
def cancel_job(job_id: int):
    """Cancel a running or queued job."""
    if current_user.role != Role.TEACHER.value:
        return redirect(url_for("teacher.dashboard"))
    job = db.session.get(ScrapeJob, job_id)
    if not job or job.teacher_id != current_user.id:
        flash("Задача не найдена.", "danger")
        return redirect(url_for("teacher.dashboard"))
    
    if job.status not in (ScrapeJobStatus.QUEUED.value, ScrapeJobStatus.RUNNING.value):
        flash("Можно отменить только задачи в статусе 'queued' или 'running'.", "warning")
        return redirect(url_for("teacher.dashboard"))
    
    # Try to kill the running process if it exists
    from ..scraper_runner import kill_job_process
    
    if kill_job_process(job_id):
        current_app.logger.info(f"Process for job {job_id} terminated successfully")
    else:
        current_app.logger.info(f"No running process found for job {job_id}")
    
    job.status = ScrapeJobStatus.CANCELLED.value
    job.finished_at = datetime.utcnow()
    if not job.error:
        job.error = "Отменено пользователем"
    db.session.commit()
    
    # IMPORTANT: Do NOT delete output directory when cancelling
    # Keep it for debugging and potential recovery
    if job.output_dir:
        current_app.logger.info(f"Job {job_id} cancelled - output directory preserved: {job.output_dir}")
    
    flash(f"Задача #{job.id} отменена.", "success")
    return redirect(url_for("teacher.dashboard"))


@bp.post("/jobs/<int:job_id>/select-school")
@login_required
def select_school_for_job(job_id: int):
    """Record user's school selection for a running job."""
    if current_user.role != Role.TEACHER.value:
        return jsonify({"error": "Unauthorized"}), 403
    
    job = db.session.get(ScrapeJob, job_id)
    if not job or job.teacher_id != current_user.id:
        return jsonify({"error": "Job not found"}), 404
    
    if job.status != ScrapeJobStatus.RUNNING.value:
        return jsonify({"error": "Job is not running"}), 400
    
    try:
        data = request.get_json()
        school_index = data.get("school_index")
        
        if school_index is None:
            return jsonify({"error": "school_index is required"}), 400
        
        # Write school selection to file for scraper to pick up
        if job.output_dir:
            from pathlib import Path
            choice_file = Path(job.output_dir) / "school_choice.txt"
            choice_file.write_text(str(school_index), encoding="utf-8")
            current_app.logger.info(f"School selection {school_index} written for job {job_id}")
            return jsonify({"success": True, "school_index": school_index})
        else:
            return jsonify({"error": "Job output directory not found"}), 500
            
    except Exception as e:
        current_app.logger.error(f"Error recording school selection for job {job_id}: {e}")
        return jsonify({"error": str(e)}), 500


@bp.post("/files/delete-all")
@login_required
def delete_all_files():
    """Delete all report files and related jobs for the current teacher."""
    if current_user.role != Role.TEACHER.value:
        flash("Доступ запрещен.", "danger")
        return redirect(url_for("teacher.dashboard"))
    
    files = ReportFile.query.filter_by(teacher_id=current_user.id).all()
    deleted_files_count = 0
    
    # Delete all report files and their physical files
    for rf in files:
        # Delete physical files if they exist
        if rf.excel_path:
            excel_path = Path(rf.excel_path)
            if excel_path.exists():
                try:
                    excel_path.unlink()
                except Exception:
                    pass  # Continue even if file deletion fails
        
        if rf.word_path:
            word_path = Path(rf.word_path)
            if word_path.exists():
                try:
                    word_path.unlink()
                except Exception:
                    pass  # Continue even if file deletion fails
        
        db.session.delete(rf)
        deleted_files_count += 1
    
    # Delete all related scrape jobs for this teacher
    jobs = ScrapeJob.query.filter_by(teacher_id=current_user.id).all()
    deleted_jobs_count = 0
    
    for job in jobs:
        # Delete output directory if it exists
        if job.output_dir:
            output_dir = Path(job.output_dir)
            if output_dir.exists():
                try:
                    shutil.rmtree(output_dir, ignore_errors=True)
                except Exception:
                    pass  # Continue even if directory deletion fails
        
        db.session.delete(job)
        deleted_jobs_count += 1
    
    # Also reset quota usage for this teacher
    quota_usages = TeacherQuotaUsage.query.filter_by(teacher_id=current_user.id).all()
    for quota in quota_usages:
        db.session.delete(quota)
    
    db.session.commit()
    
    flash(
        f"Удалено отчетов: {deleted_files_count}, "
        f"работ: {deleted_jobs_count}, "
        f"квот: {len(quota_usages)}.", 
        "success"
    )
    return redirect(url_for("teacher.dashboard"))


@bp.post("/goals/apply")
@login_required
def apply_goals():
    """Apply goals to selected reports."""
    if current_user.role != Role.TEACHER.value:
        return jsonify({"success": False, "error": "Доступ запрещен"}), 403
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "Нет данных"}), 400
        
        report_ids = data.get("report_ids", [])
        if not report_ids:
            return jsonify({"success": False, "error": "Не выбраны отчеты"}), 400
        
        # Verify all reports belong to current user
        reports = ReportFile.query.filter(
            ReportFile.id.in_(report_ids),
            ReportFile.teacher_id == current_user.id
        ).all()
        
        if len(reports) != len(report_ids):
            return jsonify({"success": False, "error": "Некоторые отчеты не найдены или недоступны"}), 403
        
        # Extract goals data
        goals_data = {
            "sor1": data.get("sor1", {}),
            "sor2": data.get("sor2", {}),
            "sor3": data.get("sor3", {}),
            "soch": data.get("soch", {}),
        }
        
        # Apply goals to Word documents
        updated_count = 0
        for rf in reports:
            if not rf.word_path:
                continue
            
            word_path = Path(rf.word_path)
            if not word_path.exists():
                continue
            
            try:
                # Import here to avoid circular imports
                import sys
                from pathlib import Path as PathLib
                # Add project root to path if needed
                project_root = PathLib(__file__).parent.parent.parent
                if str(project_root) not in sys.path:
                    sys.path.insert(0, str(project_root))
                
                from build_word_report import _fill_goals_table, _fill_difficulties_table
                
                # Open Word document
                from docx import Document
                doc = Document(str(word_path))
                
                # Fill goals table (first table)
                goals_filled = _fill_goals_table(doc, goals_data)
                
                # Fill difficulties table (third table)
                difficulties_filled = _fill_difficulties_table(doc, goals_data)
                
                if goals_filled or difficulties_filled:
                    # Save document
                    doc.save(str(word_path))
                    updated_count += 1
                    current_app.logger.info(
                        f"Updated goals in {word_path} "
                        f"(goals: {goals_filled}, difficulties: {difficulties_filled})"
                    )
                else:
                    # Log table structure for debugging
                    if current_app.debug:
                        table_info = []
                        for idx, tbl in enumerate(doc.tables):
                            if tbl.rows:
                                first_row = [c.text.strip() for c in tbl.rows[0].cells[:3]]
                                table_info.append(f"Table {idx}: {first_row}")
                        current_app.logger.warning(
                            f"Goals table not found in {word_path}. "
                            f"Found {len(doc.tables)} tables. First rows: {table_info}"
                        )
                    else:
                        current_app.logger.warning(f"Goals table not found in {word_path}")
            except Exception as e:
                current_app.logger.error(f"Error updating goals in {rf.word_path}: {e}", exc_info=True)
        
        if updated_count > 0:
            return jsonify({
                "success": True, 
                "message": f"Цели применены к {updated_count} из {len(reports)} отчетам"
            })
        else:
            return jsonify({
                "success": False, 
                "error": "Не удалось применить цели к документам"
            })
    
    except Exception as e:
        current_app.logger.error(f"Error applying goals: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


@bp.post("/goals/generate")
@login_required
def generate_goals_analysis():
    """Generate analysis text using Qwen AI based on achieved goals and difficulties."""
    if current_user.role != Role.TEACHER.value:
        return jsonify({"success": False, "error": "Доступ запрещен"}), 403
    
    # Rate Limiting проверка
    allowed, remaining = _check_rate_limit(current_user.id)
    if not allowed:
        reset_time = _get_rate_limit_reset_time(current_user.id)
        current_app.logger.warning(
            f"Rate limit exceeded for user {current_user.id}. Reset in {reset_time}s"
        )
        return jsonify({
            "success": False,
            "error": f"Превышен лимит запросов. Попробуйте через {reset_time} сек.",
            "rate_limited": True,
            "reset_in": reset_time
        }), 429
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "Нет данных"}), 400
        
        achieved = data.get("achieved", "").strip()
        difficulties = data.get("difficulties", "").strip()
        
        if not difficulties:
            return jsonify({"success": False, "error": "Заполните поле 'Цели, вызвавшие затруднения'"}), 400
        
        # Ключ и модель из школы (выбор супер-админа); fallback на env/файл для обратной совместимости
        school = current_user.school
        api_key = None
        ai_model = "qwen-flash-character"
        if school and school.ai_api_key:
            api_key = school.ai_api_key
            if school.ai_model:
                ai_model = school.ai_model
        if not api_key or api_key == "sk-xxx":
            api_key = os.getenv("DASHSCOPE_API_KEY")
        if not api_key or api_key == "sk-xxx":
            api_key_file = Path(__file__).parent.parent.parent / "api_key.txt"
            if api_key_file.exists():
                api_key = api_key_file.read_text(encoding="utf-8").strip()
        
        if not api_key or api_key == "sk-xxx":
            return _generate_fallback_analysis(achieved, difficulties)
        
        result = _call_ai_api_with_retry(api_key, achieved, difficulties, model=ai_model)
        
        if result:
            return jsonify({
                "success": True,
                "difficulties_list": result.get("difficulties_list", ""),
                "reasons": result.get("reasons", ""),
                "correction": result.get("correction", ""),
                "remaining_requests": remaining,
            })
        else:
            # Fallback если все retry не удались
            return _generate_fallback_analysis(achieved, difficulties)
    
    except Exception as e:
        current_app.logger.error(f"Error generating goals analysis: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


def _call_ai_api_with_retry(
    api_key: str,
    achieved: str,
    difficulties: str,
    model: str = "qwen-flash-character",
    max_retries: int = 3,
    base_delay: float = 1.0,
) -> dict | None:
    """
    Вызов Qwen API с retry логикой (экспоненциальная задержка).
    Модель задаётся супер-админом в кабинете школы.
    
    Args:
        api_key: API ключ
        achieved: Достигнутые цели
        difficulties: Цели с затруднениями
        model: Модель AI (qwen-flash-character, qwen-plus и т.д.)
        max_retries: Максимум попыток (default: 3)
        base_delay: Базовая задержка в секундах (default: 1.0)
    
    Returns:
        dict с результатами или None при неудаче
    """
    import json
    
    try:
        from openai import OpenAI, APIError, APIConnectionError, RateLimitError
    except ImportError:
        current_app.logger.warning("OpenAI library not installed")
        return None
    
    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
        timeout=30.0,  # Таймаут 30 секунд
    )
    
    system_prompt = """Ты помощник учителя в Казахстане. Генерируй анализ суммативного оценивания (15-30 слов на поле).

СТРОГО РАЗЛИЧАЙ три поля:

"difficulties_list" = ЧТО НЕ ПОЛУЧИЛОСЬ (какие темы/задания вызвали трудности)
Пример: "Учащиеся допускали ошибки при решении уравнений с дробями и построении графиков линейных функций."

"reasons" = ПОЧЕМУ НЕ ПОЛУЧИЛОСЬ (причины затруднений)
Пример: "Слабо усвоены правила работы с дробями, недостаточно практики в построении координатных систем."

"correction" = ЧТО ДЕЛАТЬ (план коррекционной работы)
Пример: "Провести повторение темы 'Дроби', выполнить тренировочные упражнения по построению графиков."

JSON формат: {"difficulties_list": "...", "reasons": "...", "correction": "..."}"""

    user_prompt = f"""Цели обучения:

Достигнутые: {achieved or 'Не указаны'}

С затруднениями: {difficulties}

Заполни JSON:
- difficulties_list: перечисли ЧТО не получилось
- reasons: объясни ПОЧЕМУ не получилось  
- correction: напиши ЧТО ДЕЛАТЬ для исправления"""

    last_error = None
    
    for attempt in range(max_retries):
        try:
            current_app.logger.info(f"AI API attempt {attempt + 1}/{max_retries}")
            
            completion = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                response_format={"type": "json_object"},
                temperature=0.7,
                max_tokens=500,
            )
            
            response_text = completion.choices[0].message.content
            current_app.logger.info(f"Qwen API response (attempt {attempt + 1}): {response_text}")
            
            result = json.loads(response_text)
            return result
            
        except RateLimitError as e:
            # Rate limit от API провайдера - ждем дольше
            last_error = e
            delay = base_delay * (2 ** attempt) * 2  # Удвоенная задержка
            current_app.logger.warning(
                f"API rate limit (attempt {attempt + 1}): {e}. Waiting {delay}s..."
            )
            if attempt < max_retries - 1:
                time.sleep(delay)
                
        except APIConnectionError as e:
            # Ошибка соединения - retry с задержкой
            last_error = e
            delay = base_delay * (2 ** attempt)
            current_app.logger.warning(
                f"API connection error (attempt {attempt + 1}): {e}. Waiting {delay}s..."
            )
            if attempt < max_retries - 1:
                time.sleep(delay)
                
        except APIError as e:
            # Общая ошибка API
            last_error = e
            delay = base_delay * (2 ** attempt)
            current_app.logger.warning(
                f"API error (attempt {attempt + 1}): {e}. Waiting {delay}s..."
            )
            if attempt < max_retries - 1:
                time.sleep(delay)
                
        except json.JSONDecodeError as e:
            # Ошибка парсинга JSON - не retry, сразу fallback
            current_app.logger.error(f"JSON decode error: {e}")
            return None
            
        except Exception as e:
            # Неизвестная ошибка
            last_error = e
            current_app.logger.error(f"Unexpected error (attempt {attempt + 1}): {e}")
            if attempt < max_retries - 1:
                time.sleep(base_delay)
    
    current_app.logger.error(f"All {max_retries} API attempts failed. Last error: {last_error}")
    return None


def _generate_fallback_analysis(achieved: str, difficulties: str) -> dict:
    """Fallback template-based generation when API is not available."""
    # Extract key points from difficulties
    diff_lines = [l.strip() for l in difficulties.split('\n') if l.strip()][:3]
    diff_summary = '; '.join([l.lstrip('0123456789.-* ') for l in diff_lines]).lower()
    
    difficulties_list = ""
    if diff_summary:
        difficulties_list = f"Обучающиеся испытывали затруднения при выполнении заданий по темам: {diff_summary}."
    
    reasons = "Недостаточно сформированы навыки применения теоретических знаний на практике. Требуется дополнительная работа над закреплением материала."
    
    correction = "Провести индивидуальные консультации для устранения пробелов в знаниях. Повторить теоретический материал и выполнить практические упражнения."
    
    return jsonify({
        "success": True,
        "difficulties_list": difficulties_list,
        "reasons": reasons,
        "correction": correction,
        "fallback": True,  # Indicator that fallback was used
    })

